from __future__ import annotations

from dataclasses import dataclass
import hashlib
import html
import json
from pathlib import Path
import re
from typing import Iterable


TEXT_FILE_EXTENSIONS = {
    ".html",
    ".htm",
    ".ipynb",
    ".markdown",
    ".md",
    ".mdx",
    ".pdf",
    ".rst",
    ".txt",
    ".docx",
}
EXTRACTED_TEXT_EXTENSIONS = {
    ".markdown",
    ".md",
    ".mdx",
    ".rst",
    ".txt",
}
SKIP_DIRECTORY_NAMES = {
    ".git",
    ".hg",
    ".svn",
    ".venv",
    "__pycache__",
    "node_modules",
    "raw",
}
_HTML_MAIN_PATTERNS = (
    r"<main\b[^>]*>(?P<content>[\s\S]*?)</main>",
    r"<article\b[^>]*>(?P<content>[\s\S]*?)</article>",
    r"<body\b[^>]*>(?P<content>[\s\S]*?)</body>",
)
_HTML_REMOVE_BLOCKS = ("script", "style", "svg", "noscript", "header", "footer", "nav", "aside")
_NOISE_LINE_PATTERNS = (
    r"^\s*(table of contents|contents|search|home|next|previous|up|download\s+notebook)\s*$",
    r"^\s*(edit on github|view page source|open in colab|launch binder)\s*$",
    r"^\s*(built with .*|last updated .*|copyright .*|all rights reserved.*)\s*$",
    r"^\s*(navigation|index|modules?|theme)\s*$",
    r"^\s*[-|]+\s*$",
    r"^\s*python\s+.*documentation.*\s*$",
    r"^\s*the python tutorial\s*.*\s*$",
    r"^\s*(comment|article tags:?|explore)\s*$",
)


@dataclass(frozen=True)
class CourseContentDocument:
    doc_id: str
    title: str
    relative_path: str
    source_name: str
    source_url: str | None
    text: str


def extract_text_from_file(file_path: Path) -> str:
    ext = file_path.suffix.lower()

    if ext not in TEXT_FILE_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {ext}")

    if ext in EXTRACTED_TEXT_EXTENSIONS:
        return _clean_text(file_path.read_text(encoding="utf-8", errors="replace"))

    if ext in {".html", ".htm"}:
        raw_html = file_path.read_text(encoding="utf-8", errors="replace")
        return _html_to_text(raw_html)

    if ext == ".ipynb":
        raw_notebook = file_path.read_text(encoding="utf-8", errors="replace")
        return _notebook_to_text(raw_notebook)

    if ext == ".pdf":
        try:
            import fitz
        except ImportError as exc:  # pragma: no cover - dependency install issue
            raise RuntimeError("PyMuPDF (fitz) is required for PDF parsing") from exc

        document = fitz.open(str(file_path))
        try:
            pages = [page.get_text() for page in document]
        finally:
            document.close()
        return _clean_text("\n\n".join(pages))

    if ext == ".docx":
        try:
            import docx
        except ImportError as exc:  # pragma: no cover - dependency install issue
            raise RuntimeError("python-docx is required for DOCX parsing") from exc

        document = docx.Document(str(file_path))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return _clean_text("\n\n".join(paragraphs))

    raise ValueError(f"Unsupported file type: {ext}")


def load_course_documents(course_root: Path) -> list[CourseContentDocument]:
    if not course_root.exists():
        return []

    if course_root.is_file():
        text = extract_text_from_file(course_root)
        if not text.strip():
            return []
        return [
            CourseContentDocument(
                doc_id=_stable_hash(course_root.name),
                title=_extract_title(text, course_root.stem),
                relative_path=course_root.name,
                source_name=course_root.stem.replace("_", " ").title(),
                source_url=None,
                text=text,
            )
        ]

    extracted_root = course_root / "extracted"
    if extracted_root.exists():
        files = _iter_supported_files(extracted_root)
        relative_base = extracted_root
    else:
        files = _iter_supported_files(course_root)
        relative_base = course_root

    documents: list[CourseContentDocument] = []
    for file_path in files:
        try:
            text = extract_text_from_file(file_path)
        except Exception:
            continue

        if len(text.strip()) < 120:
            continue

        relative_path = file_path.relative_to(relative_base).as_posix()
        source_name = relative_path.split("/", 1)[0].replace("-", " ").replace("_", " ").title()
        title = _extract_title(text, file_path.stem.replace("_", " ").replace("-", " ").title())
        text = _trim_navigation_preamble(text, title)
        title = _extract_title(text, title)
        documents.append(
            CourseContentDocument(
                doc_id=_stable_hash(relative_path),
                title=title,
                relative_path=relative_path,
                source_name=_extract_metadata_value(text, "Source") or source_name,
                source_url=_extract_metadata_value(text, "Original URL"),
                text=text,
            )
        )

    documents.sort(key=lambda document: document.relative_path)
    return documents


def _iter_supported_files(root: Path) -> Iterable[Path]:
    for file_path in root.rglob("*"):
        if not file_path.is_file():
            continue
        if any(part in SKIP_DIRECTORY_NAMES for part in file_path.relative_to(root).parts[:-1]):
            continue
        if file_path.suffix.lower() not in TEXT_FILE_EXTENSIONS:
            continue
        yield file_path


def _notebook_to_text(raw_notebook: str) -> str:
    try:
        data = json.loads(raw_notebook)
    except json.JSONDecodeError as exc:
        raise ValueError("Invalid notebook JSON") from exc

    parts: list[str] = []
    for cell in data.get("cells", []):
        cell_type = cell.get("cell_type")
        source = "".join(cell.get("source", [])) if isinstance(cell.get("source"), list) else str(cell.get("source", ""))
        source = source.strip()
        if not source:
            continue

        if cell_type == "markdown":
            parts.append(source)
            continue

        if cell_type == "code":
            parts.append(f"```python\n{source}\n```")
            outputs = _render_notebook_outputs(cell.get("outputs", []))
            if outputs:
                parts.append(f"Output:\n```text\n{outputs}\n```")

    return _clean_text("\n\n".join(parts))


def _render_notebook_outputs(outputs: list[dict]) -> str:
    rendered: list[str] = []
    for output in outputs:
        text = output.get("text")
        if isinstance(text, list):
            chunk = "".join(text).strip()
        elif isinstance(text, str):
            chunk = text.strip()
        else:
            data = output.get("data") or {}
            text_plain = data.get("text/plain")
            if isinstance(text_plain, list):
                chunk = "".join(text_plain).strip()
            elif isinstance(text_plain, str):
                chunk = text_plain.strip()
            else:
                chunk = ""
        if chunk:
            rendered.append(chunk[:1500])
    return "\n\n".join(rendered[:3])


def _html_to_text(raw_html: str) -> str:
    content = raw_html
    for pattern in _HTML_MAIN_PATTERNS:
        match = re.search(pattern, raw_html, flags=re.IGNORECASE)
        if match:
            content = match.group("content")
            break

    content = re.sub(r"<!--[\s\S]*?-->", " ", content)
    for tag_name in _HTML_REMOVE_BLOCKS:
        content = re.sub(
            rf"<{tag_name}\b[^>]*>[\s\S]*?</{tag_name}>",
            " ",
            content,
            flags=re.IGNORECASE,
        )

    replacements = (
        (r"<br\s*/?>", "\n"),
        (r"<li\b[^>]*>", "\n- "),
        (r"</li>", "\n"),
        (r"<(h[1-6]|p|div|section|article|main|pre|code|blockquote|tr)\b[^>]*>", "\n"),
        (r"</(h[1-6]|p|div|section|article|main|pre|code|blockquote|tr)>", "\n"),
        (r"<td\b[^>]*>", " "),
        (r"</td>", " "),
        (r"<th\b[^>]*>", " "),
        (r"</th>", " "),
    )
    for pattern, replacement in replacements:
        content = re.sub(pattern, replacement, content, flags=re.IGNORECASE)

    content = re.sub(r"<[^>]+>", " ", content)
    content = html.unescape(content)
    return _clean_text(content)


def _clean_text(text: str) -> str:
    normalized = text.replace("\r\n", "\n").replace("\r", "\n")
    normalized = normalized.replace("\u00a0", " ")
    normalized = normalized.replace("Â¶", "")
    normalized = normalized.replace("Â»", ">")

    cleaned_lines: list[str] = []
    previous_line = ""
    for raw_line in normalized.split("\n"):
        line = re.sub(r"[ \t]+", " ", raw_line).strip()
        if not line:
            if cleaned_lines and cleaned_lines[-1] != "":
                cleaned_lines.append("")
            continue
        if any(re.match(pattern, line, flags=re.IGNORECASE) for pattern in _NOISE_LINE_PATTERNS):
            continue
        if "documentation »" in line.lower():
            continue
        if line.lower() in {"comment", "article tags:", "explore"}:
            break
        if line == previous_line:
            continue
        cleaned_lines.append(line)
        previous_line = line

    text_value = "\n".join(cleaned_lines)
    text_value = re.sub(r"\n{3,}", "\n\n", text_value)
    return text_value.strip()


def _extract_metadata_value(text: str, label: str) -> str | None:
    pattern = rf"^{re.escape(label)}:\s*(.+)$"
    match = re.search(pattern, text, flags=re.IGNORECASE | re.MULTILINE)
    if not match:
        return None
    value = match.group(1).strip()
    return value or None


def _extract_title(text: str, fallback: str) -> str:
    for line in text.splitlines():
        candidate = line.strip()
        if not candidate:
            continue
        if candidate.startswith("#"):
            return candidate.lstrip("# ").strip() or fallback
        lowered = candidate.lower()
        if lowered.startswith("source:") or lowered.startswith("original url:"):
            continue
        if len(candidate) <= 120:
            return candidate
    return fallback


def _stable_hash(value: str) -> str:
    return hashlib.sha256(value.encode("utf-8")).hexdigest()[:16]


def _trim_navigation_preamble(text: str, title: str) -> str:
    raw_lines = text.splitlines()
    lines = [line.strip() for line in raw_lines]
    normalized_title = _normalize_title_text(title)
    if not normalized_title:
        return text

    metadata_end = 0
    for index, line in enumerate(lines[:10]):
        if line.lower().startswith(("source:", "original url:", "original path:", "course:")):
            metadata_end = index

    repeated_title_index = None
    short_lines_before = 0
    for index, line in enumerate(lines[metadata_end + 1 :], start=metadata_end + 1):
        normalized_line = _normalize_title_text(line)
        if normalized_line == normalized_title and index > 5:
            repeated_title_index = index
            break
        if line and len(line) <= 28:
            short_lines_before += 1

    if repeated_title_index is not None and short_lines_before >= 4:
        preserved_prefix = raw_lines[: metadata_end + 1]
        trimmed_body = raw_lines[repeated_title_index:]
        return "\n".join(preserved_prefix + [""] + trimmed_body).strip()

    return text


def _normalize_title_text(value: str) -> str:
    return re.sub(r"\s+", " ", value.replace("#", " ").replace("-", " ")).strip().lower()
